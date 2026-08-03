"""多格式知识文档加载、结构化切块与去重。"""
import hashlib
import json
import re
from pathlib import Path

from email_agent.domain.models import KnowledgeChunk


class KnowledgeBase:
    SUPPORTED_SUFFIXES = {".json", ".txt", ".md", ".docx", ".xlsx"}

    def __init__(self, knowledge_dir: str = None, chunk_size: int = 900,
                 chunk_overlap: int = 120, min_chunk_chars: int = 20):
        self.knowledge_dir = Path(knowledge_dir or Path(__file__).parent / "knowledge")
        self.knowledge_dir.mkdir(parents=True, exist_ok=True)
        self.chunk_size = max(200, int(chunk_size))
        self.chunk_overlap = min(max(0, int(chunk_overlap)), self.chunk_size // 2)
        self.min_chunk_chars = max(1, int(min_chunk_chars))
        self.entries: list[dict] = []
        self.chunks: list[KnowledgeChunk] = []
        self.load_errors: list[str] = []
        self._seen_hashes: set[str] = set()
        self._load()

    def _load(self):
        self.entries, self.chunks, self.load_errors, self._seen_hashes = [], [], [], set()
        for path in sorted(self.knowledge_dir.iterdir()):
            if not path.is_file() or path.suffix.lower() not in self.SUPPORTED_SUFFIXES:
                continue
            try:
                getattr(self, f"_load_{path.suffix.lower()[1:]}")(path)
            except Exception as exc:
                message = f"加载 {path.name} 失败: {exc}"
                self.load_errors.append(message)
                print(f"[KB Warning] {message}")

    def _load_json(self, path: Path):
        data = json.loads(path.read_text(encoding="utf-8-sig"))
        if not isinstance(data, list):
            raise ValueError("JSON 顶层必须是数组")
        for index, item in enumerate(data, 1):
            if isinstance(item, dict) and str(item.get("answer", "")).strip():
                self._add_chunks(
                    str(item.get("question", path.stem)), str(item["answer"]), path,
                    f"第 {index} 条", list(item.get("keywords", [])), {"kind": "qa"},
                )

    def _load_txt(self, path: Path):
        self._add_chunks(path.stem, path.read_text(encoding="utf-8-sig"), path, path.stem)

    def _load_md(self, path: Path):
        text = path.read_text(encoding="utf-8-sig").strip()
        text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL).strip()
        for title, content in self._markdown_sections(text, path.stem):
            self._add_chunks(title, content, path, title, metadata={"kind": "markdown"})

    def _load_docx(self, path: Path):
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("需要安装 python-docx") from exc
        document = Document(path)
        title, blocks = path.stem, []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name if paragraph.style else "").lower()
            if style.startswith("heading") or style.startswith("标题"):
                if blocks:
                    self._add_chunks(title, "\n\n".join(blocks), path, title,
                                     metadata={"kind": "docx"})
                title, blocks = text, []
            else:
                blocks.append(text)
        if blocks:
            self._add_chunks(title, "\n\n".join(blocks), path, title,
                             metadata={"kind": "docx"})
        for number, table in enumerate(document.tables, 1):
            rows = [[self._cell(cell.text) for cell in row.cells] for row in table.rows]
            rows = [row for row in rows if any(row)]
            if not rows:
                continue
            header = " | ".join(rows[0])
            for start in range(1, len(rows), 4):
                end = min(start + 4, len(rows))
                content = "\n".join([header] + [" | ".join(row) for row in rows[start:end]])
                section = f"表格 {number}（第 {start + 1}-{end} 行）"
                self._add_chunks(section, content, path, section,
                                 metadata={"kind": "table", "table": number})

    def _load_xlsx(self, path: Path):
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError("需要安装 openpyxl") from exc
        workbook = load_workbook(path, read_only=True, data_only=True)
        try:
            for sheet in workbook.worksheets:
                rows = [[self._cell(value) for value in row] for row in sheet.iter_rows(values_only=True)]
                rows = [row for row in rows if any(row)]
                if not rows:
                    continue
                header = " | ".join(rows[0])
                for start in range(1, len(rows), 4):
                    end = min(start + 4, len(rows))
                    content = "\n".join([header] + [" | ".join(row) for row in rows[start:end]])
                    section = f"{sheet.title}（第 {start + 1}-{end} 行）"
                    self._add_chunks(section, content, path, section,
                                     metadata={"kind": "xlsx", "sheet": sheet.title,
                                               "row_start": start + 1, "row_end": end})
        finally:
            workbook.close()

    def _add_chunks(self, title: str, text: str, path: Path, section: str,
                    keywords=None, metadata=None):
        pieces = self._chunks(text)
        for number, piece in enumerate(pieces, 1):
            chunk_section = section if len(pieces) == 1 else f"{section}（片段 {number}）"
            self._add(title, piece, path, chunk_section, keywords, metadata)

    def _add(self, question, answer, path: Path, section, keywords=None, metadata=None):
        answer = re.sub(r"\n{3,}", "\n\n", str(answer)).strip()
        if len(answer) < self.min_chunk_chars:
            return
        normalized = re.sub(r"\s+", " ", answer).strip().lower()
        content_hash = hashlib.sha256(normalized.encode("utf-8")).hexdigest()
        if content_hash in self._seen_hashes:
            return
        self._seen_hashes.add(content_hash)
        identifiers = tuple(self._identifiers(f"{question} {answer}"))
        values = list(keywords or []) + list(identifiers)
        values = list(dict.fromkeys(str(value).strip() for value in values if str(value).strip()))
        source = path.name
        stable = f"{source}\x1f{section}\x1f{content_hash}"
        chunk_id = hashlib.sha256(stable.encode("utf-8")).hexdigest()[:32]
        info = {"path": str(path), "mtime_ns": path.stat().st_mtime_ns,
                "identifiers": list(identifiers), **(metadata or {})}
        self.entries.append({"chunk_id": chunk_id, "question": str(question).strip(),
                             "answer": answer, "keywords": values, "source": source,
                             "section": str(section).strip(), "content_hash": content_hash,
                             "metadata": info})
        self.chunks.append(KnowledgeChunk(chunk_id, answer, source, str(section).strip(),
                                          content_hash, identifiers, info))

    def add_entry(self, question: str, answer: str, keywords: list = None):
        synthetic = self.knowledge_dir / "manual.txt"
        synthetic.touch(exist_ok=True)
        self._add(question, answer, synthetic, question, keywords)

    def get_stats(self) -> dict:
        by_type = {}
        for entry in self.entries:
            suffix = Path(entry["source"]).suffix.lower() or "manual"
            by_type[suffix] = by_type.get(suffix, 0) + 1
        return {"entries": len(self.entries), "sources": len({e["source"] for e in self.entries}),
                "by_type": by_type, "errors": list(self.load_errors),
                "min_chars": min((len(e["answer"]) for e in self.entries), default=0),
                "max_chars": max((len(e["answer"]) for e in self.entries), default=0)}

    @staticmethod
    def _markdown_sections(text: str, default: str) -> list[tuple[str, str]]:
        sections, headings, lines, title = [], [], [], default
        for line in text.splitlines():
            match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
            if not match:
                lines.append(line)
                continue
            content = "\n".join(lines).strip()
            if content:
                sections.append((title, content))
            level = len(match.group(1))
            headings = headings[:level - 1] + [re.sub(r"[*_`#]", "", match.group(2)).strip()]
            title, lines = " > ".join(headings), [line]
        content = "\n".join(lines).strip()
        if content:
            sections.append((title, content))
        return sections or [(default, text)]

    def _chunks(self, text: str) -> list[str]:
        text = str(text).strip()
        if not text:
            return []
        blocks = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
        chunks, current = [], ""
        for block in blocks:
            units = [block[i:i + self.chunk_size] for i in range(0, len(block), self.chunk_size)]
            for unit in units:
                candidate = f"{current}\n\n{unit}".strip() if current else unit
                if current and len(candidate) > self.chunk_size:
                    chunks.append(current)
                    overlap = current[-self.chunk_overlap:] if self.chunk_overlap else ""
                    current = f"{overlap}\n\n{unit}".strip()
                    if len(current) > self.chunk_size:
                        current = unit
                else:
                    current = candidate
        if current:
            chunks.append(current)
        return chunks

    @staticmethod
    def _identifiers(text: str) -> list[str]:
        candidates = re.findall(
            r"(?i)(?<![a-z0-9])[a-z][a-z0-9]*(?:[-_/\.]?[a-z0-9]+)*(?![a-z0-9])",
            text,
        )
        values = []
        for value in candidates:
            normalized = value.lower().strip("-_/.")
            if (len(normalized) > 1 and any(char.isdigit() for char in normalized)
                    and not re.fullmatch(r"(?:v|q)\d+(?:\.\d+)*", normalized)):
                values.append(normalized)
        return list(dict.fromkeys(values))

    @staticmethod
    def _chinese_terms(text: str) -> list[str]:
        stop = {"请问", "您好", "一下", "多少", "什么", "怎么", "是否", "产品", "问题"}
        terms = []
        for sequence in re.findall(r"[\u4e00-\u9fff]{2,}", text):
            for size in (2, 3):
                terms.extend(sequence[i:i + size] for i in range(len(sequence) - size + 1))
        return list(dict.fromkeys(term for term in terms if term not in stop))

    @staticmethod
    def _cell(value) -> str:
        return "" if value is None else re.sub(r"\s+", " ", str(value)).strip()
